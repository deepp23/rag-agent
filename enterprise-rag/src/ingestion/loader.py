import re
from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.core.logger import get_logger
from src.ingestion.models import DocumentBlock, LoadedDocument

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

# FIX #1: crude heuristic to detect PDF "headings" so PDFs get a
# section_path instead of always being empty (which silently disabled
# structural splitting for every PDF).
_HEADING_LIKE_PATTERN = re.compile(
    r"^(chapter|section|part)\s+\d+", re.IGNORECASE
)


def load_document(
    file_path: str | Path,
    original_file_name: str | None = None,
) -> LoadedDocument:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {extension}")

    file_name = original_file_name or path.name

    logger.info(f"Loading document: {file_name}")

    match extension:
        case ".pdf":
            return _load_pdf(path=path, file_name=file_name)
        case ".docx":
            return _load_docx(path=path, file_name=file_name)
        case ".txt":
            return _load_txt(path=path, file_name=file_name)
        case _:
            raise ValueError(f"Unsupported file type: {extension}")


# ============================================================
# DOCX (unchanged from your original — it was already correct)
# ============================================================


def _load_docx(path: Path, file_name: str) -> LoadedDocument:
    document = DocxDocument(path)
    blocks: list[DocumentBlock] = []
    section_path: list[str] = []
    block_index = 0

    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if not text:
                continue

            style_name = item.style.name.lower() if item.style else ""
            block_type = _detect_docx_block_type(style_name)

            if block_type == "heading":
                heading_level = _get_heading_level(style_name)
                section_path = _update_section_path(
                    current_path=section_path,
                    heading=text,
                    level=heading_level,
                )

            blocks.append(
                DocumentBlock(
                    block_id=f"{file_name}__block_{block_index}",
                    block_type=block_type,
                    text=text,
                    section_path=section_path.copy(),
                    metadata={"style": style_name},
                )
            )
            block_index += 1

        elif isinstance(item, Table):
            table_text = _table_to_markdown(item)
            if not table_text:
                continue

            blocks.append(
                DocumentBlock(
                    block_id=f"{file_name}__block_{block_index}",
                    block_type="table",
                    text=table_text,
                    section_path=section_path.copy(),
                    metadata={
                        "rows": len(item.rows),
                        "columns": len(item.columns) if item.columns else 0,
                    },
                )
            )
            block_index += 1

    if not blocks:
        raise ValueError(f"No content extracted from {file_name}")

    logger.info(f"Loaded DOCX {file_name}: {len(blocks)} blocks")

    return LoadedDocument(
        file_name=file_name,
        file_type="docx",
        blocks=blocks,
        total_pages=1,
    )


def _detect_docx_block_type(style_name: str) -> str:
    if style_name.startswith("heading") or style_name in ("title", "subtitle"):
        return "heading"
    if "list" in style_name or "bullet" in style_name:
        return "list_item"
    return "paragraph"


def _get_heading_level(style_name: str) -> int:
    if style_name == "title":
        return 1
    if style_name.startswith("heading"):
        try:
            return int(style_name.split()[-1])
        except ValueError:
            return 1
    return 1


def _update_section_path(
    current_path: list[str], heading: str, level: int
) -> list[str]:
    level = max(level, 1)
    new_path = current_path[: level - 1]
    new_path.append(heading)
    return new_path


def _table_to_markdown(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        if any(cells):
            rows.append(cells)

    if not rows:
        return ""

    output = [" | ".join(rows[0]), " | ".join(["---"] * len(rows[0]))]
    for row in rows[1:]:
        output.append(" | ".join(row))

    return "\n".join(output)


# ============================================================
# PDF — now with a heading heuristic feeding section_path
# ============================================================


def _load_pdf(path: Path, file_name: str) -> LoadedDocument:
    reader = PdfReader(path)
    blocks: list[DocumentBlock] = []
    block_index = 0

    # FIX #1 (continued): track a running section_path for PDFs too,
    # updated whenever a line looks like a heading.
    section_path: list[str] = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text()

        if not text or not text.strip():
            logger.warning(f"No text on page {page_number} of {file_name}")
            continue

        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        if len(paragraphs) == 1:
            paragraphs = [
                line.strip() for line in text.splitlines() if line.strip()
            ]

        for paragraph in paragraphs:
            is_heading = _looks_like_pdf_heading(paragraph)

            if is_heading:
                section_path = _update_section_path(
                    current_path=section_path,
                    heading=paragraph,
                    level=1,  # PDFs don't expose heading depth reliably
                )

            blocks.append(
                DocumentBlock(
                    block_id=f"{file_name}__block_{block_index}",
                    block_type="heading" if is_heading else "text",
                    text=paragraph,
                    page_number=page_number,
                    section_path=section_path.copy(),
                )
            )
            block_index += 1

    if not blocks:
        raise ValueError(f"No content extracted from {file_name}")

    logger.info(f"Loaded PDF {file_name}: {len(blocks)} blocks")

    return LoadedDocument(
        file_name=file_name,
        file_type="pdf",
        blocks=blocks,
        total_pages=len(reader.pages),
    )


def _looks_like_pdf_heading(line: str) -> bool:
    """
    Crude heuristic since PDFs don't expose style info via pypdf.
    Flags a line as heading-like if it's short, has no terminal
    punctuation, and either matches "Chapter/Section N" or is
    Title Case / ALL CAPS. This is imperfect but far better than
    "no structural signal at all" for every PDF.
    """
    stripped = line.strip()

    if not stripped or len(stripped) > 80:
        return False

    if stripped.endswith((".", ",", ";", ":")):
        return False

    if _HEADING_LIKE_PATTERN.match(stripped):
        return True

    words = stripped.split()
    if not words:
        return False

    if stripped.isupper() and len(words) <= 10:
        return True

    # Title Case check: most words start with a capital letter
    capitalized = sum(1 for w in words if w[:1].isupper())
    if len(words) <= 10 and capitalized / len(words) >= 0.7:
        return True

    return False


# ============================================================
# TXT (unchanged)
# ============================================================


def _load_txt(path: Path, file_name: str) -> LoadedDocument:
    text = path.read_text(encoding="utf-8").strip()

    if not text:
        raise ValueError(f"File is empty: {file_name}")

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

    blocks = [
        DocumentBlock(
            block_id=f"{file_name}__block_{i}",
            block_type="text",
            text=paragraph,
        )
        for i, paragraph in enumerate(paragraphs)
    ]

    logger.info(f"Loaded TXT {file_name}: {len(blocks)} blocks")

    return LoadedDocument(
        file_name=file_name,
        file_type="txt",
        blocks=blocks,
        total_pages=1,
    )